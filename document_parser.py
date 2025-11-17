"""
Document Parser Module
Handles PDF and Word document extraction with OCR support
"""
import pdfplumber
import PyPDF2
from docx import Document
import re
from typing import List, Tuple, Dict, Optional
import logging

# Optional OCR dependencies
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

try:
    from pdf2image import convert_from_path
    import pytesseract
    from PIL import Image
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False
    logger.warning("OCR dependencies not available. Scanned PDFs may not be processed. Install pytesseract and pdf2image for OCR support.")


class DocumentParser:
    """Parses PDF and Word documents with OCR support for scanned documents"""
    
    def __init__(self):
        self.text_content = []
        self.is_readable = True
        self.unreadable_pages = []
    
    def parse_pdf(self, file_path: str) -> Tuple[List[str], bool, List[int]]:
        """
        Parse PDF file, attempting text extraction first, then OCR if needed
        
        Returns:
            Tuple of (text_lines, is_readable, unreadable_page_numbers)
        """
        all_lines = []
        unreadable_pages = []
        
        try:
            # Try pdfplumber first (works for text-based PDFs)
            with pdfplumber.open(file_path) as pdf:
                for page_num, page in enumerate(pdf.pages, 1):
                    page_text = page.extract_text()
                    if page_text and len(page_text.strip()) > 50:
                        lines = [line.strip() for line in page_text.split('\n') if line.strip()]
                        all_lines.extend(lines)
                        logger.info(f"Page {page_num}: Extracted {len(lines)} lines via pdfplumber")
                    else:
                        # Page might be scanned/image-based, try OCR
                        if OCR_AVAILABLE:
                            logger.warning(f"Page {page_num}: Low text content, attempting OCR...")
                            try:
                                # Convert PDF page to image for OCR
                                images = convert_from_path(file_path, first_page=page_num, last_page=page_num)
                                if images:
                                    ocr_text = pytesseract.image_to_string(images[0])
                                    if len(ocr_text.strip()) > 50:
                                        lines = [line.strip() for line in ocr_text.split('\n') if line.strip()]
                                        all_lines.extend(lines)
                                        logger.info(f"Page {page_num}: Extracted {len(lines)} lines via OCR")
                                    else:
                                        unreadable_pages.append(page_num)
                                        logger.warning(f"Page {page_num}: OCR failed, insufficient text")
                                else:
                                    unreadable_pages.append(page_num)
                            except Exception as e:
                                unreadable_pages.append(page_num)
                                logger.error(f"Page {page_num}: OCR error - {str(e)}")
                        else:
                            unreadable_pages.append(page_num)
                            logger.warning(f"Page {page_num}: Low text content and OCR not available. Install pytesseract for OCR support.")
            
            is_readable = len(unreadable_pages) == 0 and len(all_lines) > 10
            return all_lines, is_readable, unreadable_pages
            
        except Exception as e:
            logger.error(f"PDF parsing error: {str(e)}")
            # Fallback to PyPDF2
            try:
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    for page_num, page in enumerate(pdf_reader.pages, 1):
                        try:
                            page_text = page.extract_text()
                            if page_text:
                                lines = [line.strip() for line in page_text.split('\n') if line.strip()]
                                all_lines.extend(lines)
                        except:
                            unreadable_pages.append(page_num)
                is_readable = len(all_lines) > 10
                return all_lines, is_readable, unreadable_pages
            except Exception as e2:
                logger.error(f"PyPDF2 fallback error: {str(e2)}")
                return [], False, list(range(1, 10))  # Assume multiple pages failed
    
    def parse_word(self, file_path: str) -> Tuple[List[str], bool, List[int]]:
        """
        Parse Word document
        
        Returns:
            Tuple of (text_lines, is_readable, unreadable_page_numbers)
        """
        try:
            doc = Document(file_path)
            all_lines = []
            
            for para in doc.paragraphs:
                if para.text.strip():
                    all_lines.append(para.text.strip())
            
            # Also check tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                    if row_text:
                        all_lines.append(row_text)
            
            is_readable = len(all_lines) > 10
            return all_lines, is_readable, []
            
        except Exception as e:
            logger.error(f"Word document parsing error: {str(e)}")
            return [], False, []
    
    def parse_document(self, file_bytes: bytes, filename: str) -> Tuple[List[str], bool, List[int]]:
        """
        Main entry point for parsing documents
        
        Args:
            file_bytes: File content as bytes
            filename: Original filename
            
        Returns:
            Tuple of (text_lines, is_readable, unreadable_page_numbers)
        """
        import tempfile
        import os
        
        # Write to temporary file
        with tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(filename)[1]) as tmp_file:
            tmp_file.write(file_bytes)
            tmp_path = tmp_file.name
        
        try:
            if filename.lower().endswith('.pdf'):
                return self.parse_pdf(tmp_path)
            elif filename.lower().endswith(('.doc', '.docx')):
                return self.parse_word(tmp_path)
            else:
                logger.error(f"Unsupported file type: {filename}")
                return [], False, []
        finally:
            # Clean up temp file
            if os.path.exists(tmp_path):
                os.unlink(tmp_path)

